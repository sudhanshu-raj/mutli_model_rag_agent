import os
import zipfile
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pandas as pd
import json
from PIL import Image
import io
import json_functions as JC
from qa_chain import QAChain
import config as cfg

OUTPUT_DIR = cfg.OUTPUT_DIR

class WordExtractor:
    def __init__(self, docx_path,workspace):
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.doc_name = os.path.splitext(os.path.basename(docx_path))[0]
        self.output_dir = os.path.join(
            OUTPUT_DIR,workspace,os.path.basename(docx_path))
        os.makedirs(self.output_dir, exist_ok=True)
        self.images_dir = os.path.join(self.output_dir, "images")
        os.makedirs(self.images_dir, exist_ok=True)

    def extract_text(self):
        text = {}
        for i, para in enumerate(self.document.paragraphs):
            text[f"paragraph_{i+1}"] = para.text

        qa_chain = QAChain()
        content_to_ask = str(text) if len(
            str(text)) < 2000 else str(text)[:2000]
        self.doc_type = qa_chain.give_document_type(content_to_ask)
        return text

    def extract_tables(self):
        tables = {}
        for table_num, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            df = pd.DataFrame(table_data[1:], columns=table_data[0])
            tables[f"table_{table_num+1}"] = df.to_markdown(index=False)
        return tables

    def extract_images(self):
        images = {}
        docx_zip = zipfile.ZipFile(self.docx_path)

        # Extract relationships to find images
        img_count = 1
        for entry in docx_zip.filelist:
            if "media" in entry.filename:
                image_data = docx_zip.read(entry.filename)
                image_ext = os.path.splitext(entry.filename)[1]
                img_name = f"image_{img_count}{image_ext}"
                img_path = os.path.join(self.images_dir, img_name)

                # Save image
                with open(img_path, "wb") as f:
                    f.write(image_data)

                # Convert to JPEG if needed and create thumbnail
                try:
                    img = Image.open(io.BytesIO(image_data))
                    images[img_name] = {
                        "path": img_path,
                        "format": img.format,
                        "size": img.size,
                        "mode": img.mode
                    }
                    img_count += 1
                except Exception as e:
                    print(f"Error processing image: {str(e)}")

        return images

    def extract_metadata(self):
        core_props = self.document.core_properties
        return {
            "author": core_props.author,
            "created": str(core_props.created),
            "modified": str(core_props.modified),
            "title": core_props.title if core_props.title else self.doc_name,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "revision": core_props.revision,
            "timestamp": str(core_props.modified),
            "source": self.docx_path,
            "doc_name": self.doc_name,
            "full_path": self.docx_path,
            "document_type": self.doc_type
        }

    def extract_all(self, extract_tables=True):
        data = {
            "text": self.extract_text(),
            "metadata": self.extract_metadata()
        }
        
        # Only extract tables if requested
        if extract_tables:
            data["tables"] = self.extract_tables()
        else:
            data["tables"] = {}  # Empty tables dictionary
            
        # Always extract images
        data["images"] = self.extract_images()

        # Save to JSON
        json_path = os.path.join(self.output_dir, "extracted_data.json")
        with open(json_path, "w") as f:
            json.dump(data, f, indent=4)
        return json_path


if __name__ == "__main__":
    docx_path = r"C:\Users\rajsu\OneDrive\Documents\amity_projects_sem4\software_design\assignment 1\Assignment_PriceScope.docx"
    extractor = WordExtractor(docx_path)
    result_path = extractor.extract_all()
    print(f"Extraction complete. Results saved to: {result_path}")
